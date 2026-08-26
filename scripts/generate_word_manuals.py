import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    if level == 1:
        run.font.name = "Segoe UI"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(227, 85, 23) # Orange
    elif level == 2:
        run.font.name = "Segoe UI"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    elif level == 3:
        run.font.name = "Segoe UI"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
    return h

def add_callout(doc, title, text, bg_color="FFF7ED", border_color="F97316"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=160, bottom=160, left=240, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"{title}\n")
    r_title.font.name = "Segoe UI"
    r_title.font.bold = True
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(194, 65, 12)
    
    r_text = p.add_run(text)
    r_text.font.name = "Segoe UI"
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def format_table_headers(tbl, col_names):
    hdr_cells = tbl.rows[0].cells
    for i, name in enumerate(col_names):
        hdr_cells[i].text = name
        set_cell_background(hdr_cells[i], "1E293B") # Slate 900
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = "Segoe UI"
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)

def format_table_rows(tbl, rows_data):
    for r_idx, row in enumerate(rows_data):
        row_cells = tbl.add_row().cells
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=120, bottom=120, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = "Segoe UI"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(30, 41, 59)

# ==============================================================================
# 1. MANUAL DE USUARIO / CLIENTE (TIENDA ONLINE)
# ==============================================================================
def create_client_manual():
    doc = docx.Document()
    
    # Page setup
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_kicker = p_title.add_run("GUÍA DE USUARIO · TIENDA ONLINE\n")
    r_kicker.font.name = "Segoe UI"
    r_kicker.font.bold = True
    r_kicker.font.size = Pt(11)
    r_kicker.font.color.rgb = RGBColor(234, 88, 12)
    
    r_main = p_title.add_run("Manual de Pedidos y Seguimiento\nBeatsBurgers & Bowls")
    r_main.font.name = "Segoe UI"
    r_main.font.bold = True
    r_main.font.size = Pt(24)
    r_main.font.color.rgb = RGBColor(15, 23, 42)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run("Paso a paso para armar tus pedidos, personalizar platos y seguir tus entregas en tiempo real.")
    r_sub.font.name = "Segoe UI"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    add_callout(
        doc,
        "💡 Resumen Rápido",
        "Podés pedir hamburguesas y combos para recibir hoy, o armar tu plan semanal de Bowls saludables eligiendo platos para distintos días de la semana en un solo checkout con seguimiento individual."
    )
    
    # Sección 1
    add_heading_styled(doc, "1. Cómo Explorar el Menú", level=1)
    p = doc.add_paragraph()
    p.add_run("Al ingresar a la aplicación web desde tu celular o computadora verás el menú organizado de forma dinámica:\n")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    r = bp1.add_run("Categorías Principales: ")
    r.bold = True
    bp1.add_run("En la barra superior podés navegar entre Burgers, Beats Bowls, Bebidas y Acompañamientos.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    r = bp2.add_run("Menú Semanal de Bowls: ")
    r.bold = True
    bp2.add_run("Cada bowl saludable tiene asignado un día específico de elaboración fresca (Lunes a Viernes). Podés ver qué ingrediente trae cada día.")

    bp3 = doc.add_paragraph(style='List Bullet')
    r = bp3.add_run("Buscador en Vivo: ")
    r.bold = True
    bp3.add_run("Podés escribir palabras clave (ej: 'Cheddar', 'Smash', 'Pollo') para encontrar tus platos favoritos.")

    # Sección 2
    add_heading_styled(doc, "2. Cómo Personalizar tu Plato", level=1)
    p = doc.add_paragraph()
    p.add_run("Hacé click o tocá cualquier hamburguesa o bowl para abrir la pantalla de personalización:\n")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    r = bp1.add_run("Quitar Ingredientes: ")
    r.bold = True
    bp1.add_run("Destildá los ingredientes que no desees (ej: 'Sin cebolla', 'Sin pepinillos', 'Sin salsa garlic').")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    r = bp2.add_run("Sumar Extras: ")
    r.bold = True
    bp2.add_run("Agregá porciones adicionales de queso cheddar, panceta crocante, medallones extra o salsas especiales.")

    bp3 = doc.add_paragraph(style='List Bullet')
    r = bp3.add_run("Notas para la Cocina: ")
    r.bold = True
    bp3.add_run("Escribí aclaraciones puntuales como 'bien cocido' o 'salsa aparte'.")

    bp4 = doc.add_paragraph(style='List Bullet')
    r = bp4.add_run("Botón Agregar: ")
    r.bold = True
    bp4.add_run("Tocá 'Agregar al Carrito' para guardar tu configuración.")

    # Sección 3
    add_heading_styled(doc, "3. Carrito de Compras y Plan Semanal Multi-Día", level=1)
    p = doc.add_paragraph()
    p.add_run("Cuando agregás platos al carrito, el sistema organiza automáticamente tu compra:\n")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    r = bp1.add_run("Pedidos de Hoy: ")
    r.bold = True
    bp1.add_run("Si pedís hamburguesas o productos diarios, se preparan y despachan en el momento o en el turno de hoy.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    r = bp2.add_run("Plan Semanal de Bowls (Varios Días): ")
    r.bold = True
    bp2.add_run("Si elegís el bowl del Lunes, el del Miércoles y el del Viernes, el carrito los agrupará por fecha de entrega. Pagás todo junto una sola vez y el sistema genera automáticamente un pedido agendado individual para cada día.")

    # Sección 4
    add_heading_styled(doc, "4. Finalizar el Pedido (Checkout)", level=1)
    p = doc.add_paragraph()
    p.add_run("Para confirmar tu compra completá los 4 pasos rápidos en la pantalla de pago:\n")
    
    bp1 = doc.add_paragraph(style='List Number')
    r = bp1.add_run("Datos de Contacto: ")
    r.bold = True
    bp1.add_run("Ingresá tu nombre y tu número de teléfono de WhatsApp (donde recibirás el estado de tu pedido).")

    bp2 = doc.add_paragraph(style='List Number')
    r = bp2.add_run("Tipo de Entrega: ")
    r.bold = True
    bp2.add_run("Elegí 'Envío a Domicilio' (ingresando tu dirección y timbre) o 'Retiro en el Local'.")

    bp3 = doc.add_paragraph(style='List Number')
    r = bp3.add_run("Franja Horaria: ")
    r.bold = True
    bp3.add_run("Elegí el horario en el que preferís recibir tus alimentos (ej: 12:00 a 13:00 hs, 20:30 a 21:30 hs).")

    bp4 = doc.add_paragraph(style='List Number')
    r = bp4.add_run("Medio de Pago: ")
    r.bold = True
    bp4.add_run("Elegí entre Efectivo al recibir o Mercado Pago (tarjetas, transferencia QR o dinero en cuenta).")

    # Sección 5
    add_heading_styled(doc, "5. Seguimiento en Vivo (Tracking en Tiempo Real)", level=1)
    p = doc.add_paragraph()
    p.add_run("Al confirmar tu pedido serás redirigido a la pantalla de Seguimiento en Vivo (/track/ID):\n")
    
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_headers(tbl, ["Tipo de Pedido", "Etapas Visibles", "Información Clave"])
    format_table_rows(tbl, [
        ["⚡ Pedidos de Hoy", "Recibido ➔ Preparando ➔ En Reparto ➔ Entregado", "Muestra el avance en tiempo real en cocina, datos del cadete con botón para llamarlo y cronómetro de demora."],
        ["📅 Pedidos Agendados (Futuros / Bowls)", "Agendado ➔ En Cocina (Día de entrega) ➔ En camino ➔ Entregado", "Muestra un cartel destacado con la fecha exacta y franja horaria pactada, recordando que se cocinará fresco ese día."],
    ])

    p_multiday = doc.add_paragraph()
    p_multiday.paragraph_format.space_before = Pt(8)
    r = p_multiday.add_run("🌟 Selector Multi-Día: ")
    r.bold = True
    p_multiday.add_run("Si compraste viandas para varios días, en la parte superior verás botones interactivos como [ 🥗 Lun 25 ] [ 🥗 Mié 27 ] [ 🥗 Jue 28 ]. Podés tocar cada día para ver el detalle y seguimiento de esa fecha.")

    # Sección 6
    add_heading_styled(doc, "6. Mi Perfil y Puntos de Recompensa", level=1)
    p = doc.add_paragraph()
    p.add_run("Desde el botón 'Mi Perfil' podés acceder a:\n")
    bp1 = doc.add_paragraph(style='List Bullet')
    r = bp1.add_run("Puntos Acumulados: ")
    r.bold = True
    bp1.add_run("Cada pedido te otorga puntos que podés canjear por descuentos o productos gratis.")
    bp2 = doc.add_paragraph(style='List Bullet')
    r = bp2.add_run("Entregas de Hoy: ")
    r.bold = True
    bp2.add_run("Tus comandas que están en elaboración o camino hoy.")
    bp3 = doc.add_paragraph(style='List Bullet')
    r = bp3.add_run("Entregas Agendadas: ")
    r.bold = True
    bp3.add_run("Tus pedidos programados para los próximos días con fecha bien visible.")
    bp4 = doc.add_paragraph(style='List Bullet')
    r = bp4.add_run("Historial: ")
    r.bold = True
    bp4.add_run("Registro de todas tus compras pasadas y comprobantes.")

    doc.save("Manual_Cliente_BeatsBurgers.docx")
    print("[OK] Manual del Cliente guardado: Manual_Cliente_BeatsBurgers.docx")

# ==============================================================================
# 2. MANUAL OPERATIVO DEL SISTEMA (ADMIN / COCINA / PERSONAL)
# ==============================================================================
def create_admin_manual():
    doc = docx.Document()
    
    # Page setup
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_kicker = p_title.add_run("MANUAL OPERATIVO DEL SISTEMA · ADMINISTRACIÓN Y COCINA\n")
    r_kicker.font.name = "Segoe UI"
    r_kicker.font.bold = True
    r_kicker.font.size = Pt(11)
    r_kicker.font.color.rgb = RGBColor(234, 88, 12)
    
    r_main = p_title.add_run("Guía Integral de Operaciones\nBeatsBurgers & Bowls")
    r_main.font.name = "Segoe UI"
    r_main.font.bold = True
    r_main.font.size = Pt(24)
    r_main.font.color.rgb = RGBColor(15, 23, 42)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run("Manual de uso completo para cajeros, cocineros, jefes de cocina y administradores.")
    r_sub.font.name = "Segoe UI"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    add_callout(
        doc,
        "📌 Propósito del Manual",
        "Este manual explica cómo gestionar la cocina en tiempo real, asignar repartidores, consultar la agenda de producción semanal de bowls, administrar ingredientes y configurar la tienda online de manera rápida y sin complicaciones."
    )
    
    # Sección 1
    add_heading_styled(doc, "1. Acceso y Estructura del Panel", level=1)
    p = doc.add_paragraph()
    p.add_run("El panel de administración está disponible ingresando a /admin desde cualquier navegador.\n")
    
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_headers(tbl, ["Módulo", "Ruta", "Objetivo Principal"])
    format_table_rows(tbl, [
        ["⚡ Pedidos Hoy (KDS)", "/admin/live", "Pantalla táctil/monitor de cocina para despachar pedidos en tiempo real."],
        ["📅 Agenda y Calendario", "/admin/calendar", "Cronograma semanal/mensual y resumen de producción de bowls a cocinar."],
        ["🍔 Catálogo y Stock", "/admin/catalog", "Gestión de productos, combos, ingredientes y control de stock."],
        ["🖼️ Galería Multimedia", "/admin/media", "Subida de fotos de productos y videos de hasta 50 MB para el Splash."],
        ["📊 Métricas y Reportes", "/admin/metricas", "Facturación total, métodos de pago, ranking de productos y rendimiento."],
        ["📜 Historial de Pedidos", "/admin/history", "Búsqueda forense, auditoría y reimpresión de comprobantes."],
        ["⚙️ Configuración General", "/admin/settings", "Horarios, costos de delivery, impresoras térmicas y diseño de la app."],
    ])

    # Sección 2
    add_heading_styled(doc, "2. Dashboard de Cocina en Vivo (/admin/live)", level=1)
    p = doc.add_paragraph()
    p.add_run("Es la pantalla operativa más importante. Diseñada como un KDS (Kitchen Display System) para operar sin distracciones:\n")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    r = bp1.add_run("Switch Local Abierto / Cerrado: ")
    r.bold = True
    bp1.add_run("Permite abrir o pausar la tienda online con un solo toque si el local está saturado.")

    bp2 = doc.add_paragraph(style='List Bullet')
    r = bp2.add_run("Métricas Rápidas: ")
    r.bold = True
    bp2.add_run("Muestra pedidos activos, pedidos del día y dinero total recaudado en el turno.")

    bp3 = doc.add_paragraph(style='List Bullet')
    r = bp3.add_run("Alertas Sonoras (Botón de Parlante): ")
    r.bold = True
    bp3.add_run("Emite una señal acústica en el monitor cada vez que un cliente envía un pedido nuevo.")

    bp4 = doc.add_paragraph(style='List Bullet')
    r = bp4.add_run("Buscador en Tiempo Real: ")
    r.bold = True
    bp4.add_run("Filtrá al instante por nombre de cliente, teléfono, plato o número de orden.")

    bp5 = doc.add_paragraph(style='List Bullet')
    r = bp5.add_run("Botón + Pedido: ")
    r.bold = True
    bp5.add_run("Permite al cajero cargar pedidos telefónicos o de mostrador sumando puntos al cliente.")

    add_heading_styled(doc, "Flujo de las 4 Columnas Operativas (Kanban)", level=2)
    
    tbl_kanban = doc.add_table(rows=1, cols=3)
    tbl_kanban.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_headers(tbl_kanban, ["Columna", "Qué representa", "Acción requerida"])
    format_table_rows(tbl_kanban, [
        ["🔔 Nuevos / Ingresados", "Comandas recién llegadas por la web.", "Verificar notas e ingredientes y presionar 'Cocinar 🔥' para aceptar, o 'Rechazar'."],
        ["🔥 En Cocina", "Platos elaborándose en cocina.", "Al terminar: si es delivery seleccionar repartidor y presionar 'A Reparto 🛵'. Si es retiro, presionar 'Listo Retiro ✅'."],
        ["🛵 Despacho / En Reparto", "Pedidos en viaje con el cadete o esperando al cliente en mostrador.", "Presionar 'WS Cadete' para enviar la hoja de ruta con mapa al repartidor. Al entregar, presionar 'Marcar Entregado ✅'."],
        ["✅ Completados Hoy", "Historial de pedidos cerrados hoy.", "Columna compacta y plegable para no estorbar la visibilidad de la cocina."],
    ])

    add_heading_styled(doc, "Temporizadores de Demora en Comandas", level=2)
    p_temp = doc.add_paragraph()
    p_temp.add_run("Cada tarjeta de comanda calcula el tiempo transcurrido desde su ingreso:\n")
    doc.add_paragraph("• 🟢 Verde (0 a 14 min): Tiempo óptimo de preparación.", style='List Bullet')
    doc.add_paragraph("• 🟡 Amarillo (15 a 24 min): Tiempo de atención sugerido.", style='List Bullet')
    doc.add_paragraph("• 🔴 Rojo Parpadeante (25+ min): Alerta de demora crítica para priorizar en cocina.", style='List Bullet')

    # Sección 3
    add_heading_styled(doc, "3. Módulo de Agenda y Calendario (/admin/calendar)", level=1)
    p = doc.add_paragraph()
    p.add_run("Especialmente desarrollado para gestionar el Menú Semanal de Beats Bowls y pedidos por encargo:\n")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    r = bp1.add_run("Vistas de Navegación: ")
    r.bold = True
    bp1.add_run("Podés alternar entre vista de Semana (7 columnas de Lunes a Domingo), vista de Mes (grilla mensual con volumen de ventas) y vista de Día (desglose horario).")

    bp2 = doc.add_paragraph(style='List Bullet')
    r = bp2.add_run("Resumen de Producción (Prep Summary): ")
    r.bold = True
    bp2.add_run("En la parte superior calcula automáticamente la suma exacta de viandas a cocinar ese día (ej: 14x Chicken Pasta Bowl, 8x Ocean Bowl). Permite al equipo de cocina preelaborar y armar los mise en place con total exactitud.")

    # Sección 4
    add_heading_styled(doc, "4. Catálogo, Ingredientes y Stock (/admin/catalog)", level=1)
    p = doc.add_paragraph()
    p.add_run("Administración integral de la oferta gastronómica:\n")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    r = bp1.add_run("Control de Stock Automático: ")
    r.bold = True
    bp1.add_run("Cada venta descuenta las unidades de ingredientes correspondientes (medallones, panes, vegetales, salsas).")

    bp2 = doc.add_paragraph(style='List Bullet')
    r = bp2.add_run("Alerta de Falta de Stock: ")
    r.bold = True
    bp2.add_run("Si un ingrediente llega a 0, la comanda avisa con un cartel rojo y ofrece el botón 'Reponer Ahora' para restablecer las existencias al instante.")

    bp3 = doc.add_paragraph(style='List Bullet')
    r = bp3.add_run("Disponibilidad por Día: ")
    r.bold = True
    bp3.add_run("Podés configurar productos diarios (disponibles siempre) o restringidos a días específicos (ej: MONDAY para el bowl del lunes).")

    # Sección 5
    add_heading_styled(doc, "5. Galería Multimedia y Videos de Splash (/admin/media)", level=1)
    p = doc.add_paragraph()
    p.add_run("Gestión de activos visuales para la aplicación:\n")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    r = bp1.add_run("Subida de Fotos y Videos (hasta 50 MB): ")
    r.bold = True
    bp1.add_run("Permite subir imágenes (JPG, PNG, WEBP) y videos (MP4, WEBM, MOV) directamente desde tu equipo por arrastrar y soltar.")

    bp2 = doc.add_paragraph(style='List Bullet')
    r = bp2.add_run("Video de Bienvenida (Splash): ")
    r.bold = True
    bp2.add_run("En /admin/settings ➔ Diseño podés activar el Splash en modo Video y elegir cualquier archivo de la galería para que se reproduzca automáticamente sin sonido al abrir la aplicación.")

    # Sección 6
    add_heading_styled(doc, "6. Impresión Térmica y WhatsApp", level=1)
    p = doc.add_paragraph()
    p.add_run("El sistema soporta dos modalidades de impresión de tickets:\n")
    doc.add_paragraph("• Modo Navegador (Pop-up): Abre el ticket formateado en 58mm u 80mm listo para imprimir con Ctrl+P.", style='List Bullet')
    doc.add_paragraph("• Modo PrintNode: Envía el ticket de cocina y de mostrador directamente a impresoras térmicas de red o USB sin requerir confirmación manual.", style='List Bullet')
    doc.add_paragraph("• Comunicación por WhatsApp: Cada comanda y el modal de detalle cuentan con enlace directo para contactar al cliente con el mensaje preformateado.", style='List Bullet')

    doc.save("Manual_Administrador_BeatsBurgers.docx")
    print("[OK] Manual del Administrador guardado: Manual_Administrador_BeatsBurgers.docx")

if __name__ == "__main__":
    create_client_manual()
    create_admin_manual()
